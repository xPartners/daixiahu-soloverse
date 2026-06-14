"""政策文档解析引擎：将政策文件转化为结构化规则。

支持 PDF / Word / 文本类政策文件的解析，结合 RuleExtractor 抽取规则，
并可选地持久化到数据库与索引到 Elasticsearch。
"""
from __future__ import annotations

import io
import logging
import uuid
from datetime import datetime
from typing import Optional

from app.core.exceptions import PolicyParseError
from app.data.search import PolicySearch
from app.data.storage import FileStorage
from app.models.rule_extractor import RuleExtractor

logger = logging.getLogger(__name__)

# README 声称的准确率
_ACCURACY_ESTIMATE = 0.95


def _guess_content_type(source_type: str) -> str:
    """根据来源类型猜测 MIME 类型。"""
    mapping = {
        "pdf": "application/pdf",
        "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text": "text/plain",
        "web": "text/html",
        "scan": "application/octet-stream",
    }
    return mapping.get(source_type.lower(), "application/octet-stream")


class PolicyParseEngine:
    """政策文档解析引擎，将政策文件转化为结构化规则。"""

    def __init__(
        self,
        extractor: Optional[RuleExtractor] = None,
        storage: Optional[FileStorage] = None,
        search: Optional[PolicySearch] = None,
        db_session=None,
    ) -> None:
        self.extractor: RuleExtractor = extractor or RuleExtractor()
        self.storage: Optional[FileStorage] = storage
        self.search: Optional[PolicySearch] = search
        self.db_session = db_session

    async def parse_document(
        self,
        filename: str,
        region: str,
        content: bytes,
        source_type: Optional[str] = None,
    ) -> dict:
        """解析政策文档，返回结构化结果。

        source_type 为 None 时自动检测文件格式（基于文件头魔数），
        也可由用户显式指定（pdf/word/text）用于改名文件等边缘场景。

        返回 dict 含 document, rules, rule_count, accuracy_estimate。
        """
        # 0. 自动检测文件格式（用户未指定时）
        detected_type, detect_reason = self._detect_source_type(
            content, filename, source_type
        )
        logger.info(
            "文件格式判定：%s（依据：%s，文件名：%s）",
            detected_type, detect_reason, filename,
        )

        # 1. 上传原始文件（storage 可用时）
        file_url: Optional[str] = None
        if self.storage is not None:
            try:
                object_name = f"policies/{uuid.uuid4().hex}/{filename}"
                file_url = self.storage.upload(
                    object_name, content, content_type=_guess_content_type(detected_type)
                )
            except Exception as exc:  # noqa: BLE001
                logger.warning("文件上传失败，跳过: %s", exc)
                file_url = None

        # 2. 提取文本
        text = self._extract_text(content, detected_type)

        # 3. 抽取规则
        rules = self.extractor.extract(text, region=region)

        # 4. 索引到 ES（容错）
        doc_id = uuid.uuid4().hex
        if self.search is not None:
            try:
                await self.search.index_policy(
                    doc_id,
                    {
                        "title": filename,
                        "region": region,
                        "content": text,
                        "source_type": detected_type,
                        "rule_count": len(rules),
                    },
                )
            except Exception as exc:  # noqa: BLE001
                logger.warning("ES 索引失败，跳过: %s", exc)

        # 5. 持久化到 DB（db_session 可用时）
        created_at = datetime.now()
        if self.db_session is not None:
            await self._persist(doc_id, filename, region, detected_type, file_url, text, rules)

        # 6. 返回结果
        return {
            "document": {
                "id": doc_id,
                "title": filename,
                "region": region,
                "source_type": detected_type,
                "status": "parsed",
                "created_at": created_at,
                "parsed_rules_count": len(rules),
            },
            "rules": rules,
            "rule_count": len(rules),
            "accuracy_estimate": _ACCURACY_ESTIMATE,
        }

    async def _persist(
        self,
        doc_id: str,
        filename: str,
        region: str,
        source_type: str,
        file_url: Optional[str],
        text: str,
        rules: list[dict],
    ) -> None:
        """持久化 PolicyDocument 和 MedicalRule 记录（容错）。"""
        from app.data.models.policy import PolicyDocument
        from app.data.models.rule import MedicalRule

        try:
            policy_doc = PolicyDocument(
                id=doc_id,
                title=filename,
                region=region,
                source_type=source_type,
                file_url=file_url,
                raw_text=text,
                status="parsed",
                parsed_content={"rules": rules},
                rule_count=len(rules),
            )
            self.db_session.add(policy_doc)
            for rule in rules:
                medical_rule = MedicalRule(
                    rule_code=f"{doc_id}_{uuid.uuid4().hex[:8]}",
                    name=rule.get("name", ""),
                    element_type=rule.get("element_type", "benefit"),
                    region=rule.get("region", region),
                    logic=rule.get("logic", {}),
                    priority=rule.get("priority", 50),
                    status=rule.get("status", "draft"),
                    description=rule.get("description", ""),
                    action=rule.get("action", "reject"),
                    source_document_id=doc_id,
                )
                self.db_session.add(medical_rule)
            await self.db_session.commit()
        except Exception as exc:  # noqa: BLE001
            logger.warning("DB 持久化失败，跳过: %s", exc)
            try:
                await self.db_session.rollback()
            except Exception:  # noqa: BLE001
                pass

    @staticmethod
    def _detect_source_type(
        content: bytes, filename: Optional[str], hint: Optional[str]
    ) -> tuple[str, str]:
        """自动检测文件格式，返回 (source_type, 判定依据说明)。

        判定优先级：用户显式指定 > 文件头魔数 > 文本解码尝试。
        无法识别时抛出带明确归因与建议的 PolicyParseError。
        """
        # 1. 用户显式指定且合法，直接采用（用于改名文件等边缘场景）
        valid_types = {"pdf", "word", "text", "web", "scan"}
        if hint and hint.lower() in valid_types:
            return hint.lower(), f"用户手动指定（source_type={hint}）"

        # 2. 文件头魔数检测（最可靠）
        if content[:4] == b"%PDF":
            return "pdf", "文件头魔数 %PDF（判定为 PDF）"
        if content[:4] == b"PK\x03\x04":
            return "word", "文件头魔数 PK（Office OpenXML，判定为 Word .docx）"
        if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            raise PolicyParseError(
                "检测到老版 Word 文档（.doc 二进制 OLE2 格式），"
                "当前解析器(python-docx)不支持。"
                "建议：用 Word 打开后另存为 .docx 格式，再重新上传。"
            )

        # 3. 文本类：依次尝试多种编码解码
        for enc in ("utf-8", "gbk", "gb18030"):
            try:
                content.decode(enc)
                return "text", f"成功以 {enc} 编码解码（判定为纯文本）"
            except UnicodeDecodeError:
                continue

        # 4. 全部失败：输出归因与建议
        size_kb = len(content) / 1024
        raise PolicyParseError(
            f"无法自动识别文件格式（文件名：{filename}，大小：{size_kb:.1f}KB）。"
            "支持：文本型PDF、Word(.docx)、纯文本(.txt)。"
            "可能原因：文件损坏、为扫描图片型PDF(需OCR暂不支持)、"
            "或为其他二进制格式。请确认文件类型后重新上传。"
        )

    def _extract_text(self, content: bytes, source_type: str) -> str:
        """根据已判定的文件类型提取文本，并输出解析器日志。

        - pdf: 用 PyPDF2.PdfReader 读取（空文本提示扫描件）
        - word: 用 docx.Document 读取
        - text: 尝试 utf-8 / gbk / gb18030 解码
        - 异常时抛带归因的 PolicyParseError
        """
        source_type = source_type.lower()
        logger.info("调用解析器：%s", source_type)
        if source_type == "pdf":
            return self._extract_pdf(content)
        if source_type == "word":
            return self._extract_word(content)
        # text / web：尝试多种编码
        for enc in ("utf-8", "gbk", "gb18030"):
            try:
                text = content.decode(enc)
                logger.info("文本解码成功（编码：%s，字符数：%d）", enc, len(text))
                return text
            except UnicodeDecodeError:
                continue
        raise PolicyParseError(
            f"文本解码失败：utf-8 / gbk / gb18030 均无法解码。"
            "可能文件并非纯文本，或编码异常。"
        )

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """从 PDF 提取文本。若文本极少，提示可能为扫描件。"""
        try:
            import PyPDF2
        except ImportError as exc:
            raise PolicyParseError("PyPDF2 未安装，无法解析 PDF") from exc
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            texts: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
            result = "\n".join(texts).strip()
            # 启发式判定：页数>0 但提取文本极少，很可能为扫描型PDF
            if page_count > 0 and len(result) < 10:
                raise PolicyParseError(
                    f"PDF 共 {page_count} 页，但仅提取到 {len(result)} 个字符。"
                    "该文件很可能为扫描型PDF（图片型），当前系统暂不支持 OCR 文字识别。"
                    "建议：使用带可选文字的文本型PDF，或手动录入文本后以 .txt 上传。"
                )
            logger.info("PDF 解析成功（页数：%d，字符数：%d）", page_count, len(result))
            return result
        except PolicyParseError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise PolicyParseError(
                f"PDF 解析失败，归因：{exc}。"
                "可能文件已损坏或为加密PDF。"
            ) from exc

    @staticmethod
    def _extract_word(content: bytes) -> str:
        """从 Word 文档提取文本。"""
        try:
            import docx
        except ImportError as exc:
            raise PolicyParseError("python-docx 未安装，无法解析 Word 文档") from exc
        try:
            document = docx.Document(io.BytesIO(content))
            texts: list[str] = []
            for para in document.paragraphs:
                if para.text:
                    texts.append(para.text)
            return "\n".join(texts)
        except Exception as exc:  # noqa: BLE001
            raise PolicyParseError(f"Word 文档解析失败: {exc}") from exc
